from docx import Document
from typing import Dict, List


def extract_docx_content(docx_path: str) -> Dict[str, Any]:
    """
    Extrai conteúdo do arquivo DOCX
    """
    try:
        doc = Document(docx_path)

        content = {
            "paragraphs": [],
            "tables": [],
            "full_text": ""
        }

        # Extrair parágrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content["paragraphs"].append(paragraph.text)

        # Extrair tabelas
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            content["tables"].append(table_data)

        # Texto completo
        content["full_text"] = "\n".join(content["paragraphs"])

        return content

    except Exception as e:
        raise Exception(f"Erro ao processar DOCX: {str(e)}")
