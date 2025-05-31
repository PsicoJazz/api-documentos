from docx import Document
from docx.shared import Inches
import os
from typing import Dict, List, Any


def fill_template(template_path: str, xml_data: Dict[str, Any], docx_content: Dict[str, Any], request_id: str) -> str:
    """
    Preenche o template com os dados extraídos
    """
    try:
        # Abrir template
        doc = Document(template_path)

        # Dicionário de substituições
        replacements = {
            "{{CLIENTE}}": xml_data.get("cliente", ""),
            "{{DATA}}": xml_data.get("data", ""),
            # Limitar tamanho
            "{{CONTEUDO}}": docx_content.get("full_text", "")[:500],
        }

        # Adicionar valores do XML
        for key, value in xml_data.get("valores", {}).items():
            replacements[f"{{{{{key.upper()}}}}}"] = str(value)

        # Substituir texto nos parágrafos
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        placeholder, replacement)

        # Substituir texto nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, replacement in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(
                                placeholder, replacement)

        # Adicionar conteúdo adicional se necessário
        if xml_data.get("items"):
            # Adicionar parágrafo de separação
            doc.add_paragraph("Itens do XML:")

            # Adicionar itens como lista
            for item in xml_data["items"]:
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(f"Item: {item}")

        # Salvar documento final
        output_path = f"output/documento-{request_id}.docx"
        doc.save(output_path)

        return output_path

    except Exception as e:
        raise Exception(f"Erro ao preencher template: {str(e)}")
